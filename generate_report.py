from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_COLOR_INDEX

def get_paragraphs(workbook):
    worksheet = workbook.active
    paragraphs = []
    titles = []
    for row in range(2, worksheet.max_row + 1):  # Skip header row
        paragraphs.append(worksheet.cell(row, 1).value)
        titles.append(worksheet.cell(row, 2).value)
    return paragraphs, titles

def find_longest_shortest_middle(paragraphs):
    longest = max(paragraphs, key=len)
    shortest = min(paragraphs, key=len)
    middle_index = len(paragraphs) // 2
    middle = paragraphs[middle_index]
    return longest, shortest, middle

def underline_words_with_r(paragraph):
    for run in paragraph.runs:
        for word in run.text.split():
            if word.endswith('r'):
                # Create a new run with underline formatting
                new_run = paragraph.add_run(word)
                new_run.font.underline = True

def underline_words_in_document(document):
    for paragraph in document.paragraphs:
        underline_words_with_r(paragraph)

def create_report(paragraphs, titles, longest, shortest, middle):
    document = Document()

    # Add longest paragraph with "Fish" heading
    document.add_heading("Fish", level=1)
    p1 = document.add_paragraph(longest)

    # Add middle paragraph with "Cheese" heading
    document.add_heading("Cheese", level=1)
    p2 = document.add_paragraph(middle)

    # Add shortest paragraph with "Car" heading
    document.add_heading("Car", level=1)
    p3 = document.add_paragraph(shortest)

    # Add image of a python
    document.add_picture("python_image.png", width=Inches(6), height=Inches(3.375))

    return document

def main():
    workbook = load_workbook("data.xlsx")  
    paragraphs, titles = get_paragraphs(workbook)
    longest, shortest, middle = find_longest_shortest_middle(paragraphs.copy())  
    document = create_report(paragraphs, titles, longest, shortest, middle)
    
    # Underline words ending with 'r' in the entire document
    underline_words_in_document(document)
    
    document.save("report.docx")

if __name__ == "__main__":
    main()
